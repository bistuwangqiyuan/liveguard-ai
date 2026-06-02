"""
render_docx.py
==============

把 `守播LiveGuard_商业计划书_v2.0.md` 渲染为同名 .docx：

  1. 用 python-docx 生成中文 reference.docx（标题样式 / 字体 / 页眉页脚 / 封面）
  2. 用 pandoc --reference-doc=reference.docx 转换 markdown → docx

Windows 友好：CJK 字体使用 "Microsoft YaHei"（微软雅黑）。
"""

from __future__ import annotations

import subprocess
import sys
from pathlib import Path

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).parent.parent.parent
BUILD = ROOT / "liveguard" / "bp_build"
REF = BUILD / "reference.docx"
MD = ROOT / "守播LiveGuard_商业计划书_v4.0.md"
OUT = ROOT / "守播LiveGuard_商业计划书_v4.0.docx"

LATIN_FONT = "Microsoft YaHei"
CJK_FONT = "Microsoft YaHei"
HEADING_FONT = "Microsoft YaHei"

PRIMARY = "1F6FFF"
INK = "0B0F1A"
GREY = "8A93A6"


def set_cjk_font(target, name: str = CJK_FONT):
    if hasattr(target, "_element"):
        rPr = target._element.get_or_add_rPr()
    else:
        rPr = target
        if rPr is None:
            return
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:eastAsia", "w:hAnsi", "w:cs", "w:ascii"):
        rFonts.set(qn(attr), name)


def style_run(run, font_name: str = LATIN_FONT, size_pt: float = 11,
              bold: bool = False, color_hex: str = INK):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color_hex)
    set_cjk_font(run, CJK_FONT)


def build_reference():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    rPr_normal = normal.element.find(qn("w:rPr"))
    if rPr_normal is None:
        rPr_normal = OxmlElement("w:rPr")
        normal.element.append(rPr_normal)
    set_cjk_font(rPr_normal, CJK_FONT)

    for h, size, color in [
        ("Heading 1", 22, PRIMARY),
        ("Heading 2", 16, INK),
        ("Heading 3", 13, INK),
        ("Heading 4", 11, INK),
    ]:
        st = styles[h]
        st.font.name = HEADING_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        rPr_h = st.element.find(qn("w:rPr"))
        if rPr_h is None:
            rPr_h = OxmlElement("w:rPr")
            st.element.append(rPr_h)
        set_cjk_font(rPr_h, HEADING_FONT)

    header = section.header
    p = header.paragraphs[0]
    r = p.add_run("守播 LiveGuard AI · 商业计划书 v3.0")
    style_run(r, size_pt=9, color_hex=GREY)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    p = footer.paragraphs[0]
    r = p.add_run("CONFIDENTIAL · 仅授权对象阅读 · © 2026 守播 LiveGuard AI")
    style_run(r, size_pt=8.5, color_hex=GREY)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("\n\n\n\n守播 LiveGuard AI")
    style_run(r, font_name=HEADING_FONT, size_pt=36, bold=True, color_hex=PRIMARY)
    cover.add_run("\n")
    r = cover.add_run("直播经济实时可信与风控中台\n商业计划书 v3.0 · 天使 IRR 优先")
    style_run(r, font_name=HEADING_FONT, size_pt=18, color_hex=INK)
    cover.add_run("\n")
    r = cover.add_run("\n2026 年 06 月 01 日")
    style_run(r, size_pt=12, color_hex=GREY)

    doc.save(REF)
    print(f"✓ reference.docx 写入 {REF}")


def render_docx():
    print(f"── pandoc render: {MD.name} → {OUT.name} ──")
    cmd = [
        "pandoc", str(MD), "-o", str(OUT),
        "--reference-doc", str(REF),
        "--toc", "--toc-depth=2",
        "--from", "markdown+pipe_tables+task_lists",
        "--to", "docx", "--standalone",
        "--resource-path", str(ROOT),
    ]
    res = subprocess.run(cmd, capture_output=True, text=True, encoding="utf-8")
    if res.returncode != 0:
        print("STDERR:", res.stderr)
        raise SystemExit(res.returncode)
    print(f"✓ 写入 {OUT}  大小 {OUT.stat().st_size/1024:.1f} KB")


if __name__ == "__main__":
    build_reference()
    render_docx()
